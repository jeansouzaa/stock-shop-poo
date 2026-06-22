import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def parse_cs_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    namespace_match = re.search(r'namespace\s+([\w\.]+)', content)
    namespace = namespace_match.group(1) if namespace_match else "StockShop"
    classes = []
    class_pattern = r'(public|internal)?\s+(abstract\s+)?(class|interface|enum|struct)\s+(\w+)(?:\s*:\s*([\w\s,<>]+))?'
    for match in re.finditer(class_pattern, content):
        cls_type = match.group(3)
        cls_name = match.group(4)
        cls_base = match.group(5).strip() if match.group(5) else ""
        classes.append({
            'type': cls_type,
            'name': cls_name,
            'base': cls_base,
            'methods': [],
            'properties': [],
            'fields': []
        })
    if not classes:
        return None
    lines = content.split('\n')
    current_class = classes[0]
    for line in lines:
        line = line.strip()
        if not line:
            continue
        prop_match = re.match(r'(public|protected|internal|private)\s+([\w<>_\[\]\?]+)\s+(\w+)\s*\{\s*(get|set|init|private\s+set|protected\s+set)', line)
        if prop_match:
            current_class['properties'].append({
                'access': prop_match.group(1),
                'type': prop_match.group(2),
                'name': prop_match.group(3)
            })
            continue
        method_match = re.match(r'(public|protected|internal|private)\s+(?:(static|virtual|override|abstract)\s+)?([\w<>_\[\]\?]+)\s+(\w+)\s*\(([^)]*)\)', line)
        if method_match:
            access = method_match.group(1)
            modifier = method_match.group(2) or ""
            ret_type = method_match.group(3)
            name = method_match.group(4)
            args = method_match.group(5)
            if name in ('if', 'while', 'switch', 'for', 'foreach', 'catch'):
                continue
            current_class['methods'].append({
                'access': access,
                'modifier': modifier,
                'type': ret_type,
                'name': name,
                'args': args
            })
            continue
        field_match = re.match(r'(private|protected|public|internal)\s+(?:(static|const)\s+)?([\w<>_\[\]]+)\s+(\w+)\s*(?:=.*)?\s*;', line)
        if field_match:
            current_class['fields'].append({
                'access': field_match.group(1),
                'type': field_match.group(3),
                'name': field_match.group(4)
            })
            continue
    return {
        'filepath': filepath,
        'filename': os.path.basename(filepath),
        'namespace': namespace,
        'classes': classes
    }

def gather_codebase_info(root_dir):
    categories = {'Models': [], 'Controllers': [], 'Views': [], 'Other': []}
    for dirpath, _, filenames in os.walk(root_dir):
        if 'bin' in dirpath or 'obj' in dirpath or '.git' in dirpath or '.gemini' in dirpath:
            continue
        for name in filenames:
            if name.endswith('.cs'):
                full_path = os.path.join(dirpath, name)
                info = parse_cs_file(full_path)
                if not info:
                    continue
                rel_path = os.path.relpath(full_path, root_dir)
                parts = rel_path.split(os.sep)
                if len(parts) > 1 and parts[0] in categories:
                    categories[parts[0]].append(info)
                else:
                    categories['Other'].append(info)
    return categories

def build_docx(output_path, root_dir):
    code_info = gather_codebase_info(root_dir)
    doc = docx.Document()

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Documentação Técnica: Sistema de Controle de Estoque StockShop")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)

    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 65, 85)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_body_text(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.font.name = 'Arial'
            run_prefix.font.size = Pt(9.5)
            run_prefix.font.bold = True
            run_prefix.font.color.rgb = RGBColor(51, 65, 85)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(15, 23, 42)
        return p

    add_heading_1("1. Documento de Requisitos do Produto (PRD)")
    add_body_text(" Sistema StockShop", "Nome do Produto:")
    add_body_text(" Lojistas, gerentes de estoque e administradores de comércios de pequeno porte.", "Público-Alvo:")
    add_body_text(" Fornecer uma interface de linha de comando (CLI) simples e eficiente para o gerenciamento de produtos, fornecedores e movimentações de estoque, permitindo operações de cadastro, consulta, alteração e exclusão (CRUD), além da geração de relatórios consolidados de controle.", "Objetivo:")
    
    add_body_text("O sistema foi desenvolvido em C# utilizando princípios de Programação Orientada a Objetos (POO) e padrão arquitetural MVC. A interface interativa de console é desenhada com caracteres ASCII, apresentando cores personalizadas e controle de cursor. A persistência de dados é mantida inteiramente em memória por meio de estruturas estáticas durante a execução.", "Visão Geral:")
    
    add_body_text("", "Escopo:")
    add_bullet("Gerenciamento de Fornecedores (Cadastro, Consulta, Alteração e Exclusão com CNPJ único).")
    add_bullet("Gerenciamento de Produtos (Cadastro vinculando Fornecedor, Preço, Categoria de 0 a 5 e Estoque).")
    add_bullet("Registro de Movimentações de Estoque (Lançamentos de entrada e saída com controle automático de saldo físico).")
    add_bullet("Relatórios Consolidados (Listagens tabulares estruturadas para produtos, fornecedores e movimentações).")

    add_heading_1("2. Regras de Negócio")
    add_body_text("Produtos são identificados exclusivamente por um Código numérico.", "1. Identificação Única dos Produtos:")
    add_body_text("Fornecedores são identificados por Código e por um CNPJ estruturado.", "2. Identificação Única dos Fornecedores:")
    add_body_text("Movimentações de estoque possuem um Código de identificação exclusivo.", "3. Identificação Única das Movimentações:")
    
    add_body_text("Não é permitido o cadastro de dois fornecedores com o mesmo CNPJ.", "4. Validação de CNPJ Único:")
    add_body_text("Não é permitido o cadastro de dois produtos com o mesmo Código.", "5. Validação de Código Único de Produto:")
    add_body_text("Para registrar um produto, é obrigatório informar um fornecedor existente no sistema. Caso o fornecedor não exista, o fluxo disponibiliza a opção de cadastrá-lo imediatamente.", "6. Relacionamento de Cadastro de Produto:")
    add_body_text("Para registrar uma movimentação de estoque, o produto selecionado deve estar cadastrado.", "7. Relacionamento de Movimentação:")
    
    add_body_text("Toda movimentação de produto deduz fisicamente a quantidade do estoque do produto. O sistema impede movimentações cujo saldo resultante seja negativo, mostrando erro de estoque insuficiente e reiniciando a entrada.", "8. Controle de Saldo de Estoque:")
    add_body_text("Na exclusão de uma movimentação, a quantidade movimentada é estornada (somada de volta) ao estoque do produto.", "9. Estorno na Exclusão de Movimentação:")
    add_body_text("Na alteração, o sistema devolve o estoque anterior temporariamente para validar o novo lançamento contra o estoque atualizado.", "10. Estorno na Alteração de Movimentação:")
    add_body_text("Os dados são mantidos em coleções estáticas em memória (Listas). Encerrando o aplicativo, os dados são resetados.", "11. Persistência dos Dados:")

    add_heading_1("3. Requisitos Funcionais")
    
    requirements = [
        ("RF01", "Gerenciar Produtos", "Permite cadastrar, consultar, alterar e excluir produtos. Os dados incluem código, descrição, categoria (enum), preço, estoque e fornecedor."),
        ("RF02", "Gerenciar Fornecedores", "Permite cadastrar, consultar, alterar e excluir fornecedores. O CNPJ é validado contra duplicidade no sistema."),
        ("RF03", "Registrar Movimentação", "Registra movimentações vinculando código do produto, data e quantidade, validando e subtraindo a quantidade física do estoque."),
        ("RF04", "Consultar Movimentação", "Permite consultar lançamentos anteriores informando o código da movimentação."),
        ("RF05", "Excluir Movimentação", "Remove uma movimentação cadastrada e estorna a quantidade movimentada de volta ao estoque do produto correspondente."),
        ("RF06", "Relatório de Produtos", "Exibe tabela formatada com todos os produtos, categorias, preços, estoque atual e seus respectivos fornecedores."),
        ("RF07", "Relatório de Fornecedores", "Exibe tabela formatada com código, CNPJ, nome e telefone de todos os fornecedores."),
        ("RF08", "Relatório de Movimentações", "Exibe histórico de movimentações efetuadas no estoque."),
        ("RF09", "Interface CLI Dinâmica", "Apresenta menu interativo com molduras ASCII e tratamento resiliente para execução sem console físico.")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Nome'
    hdr_cells[2].text = 'Descrição'
    for cell in hdr_cells:
        set_cell_background(cell, "0F172A")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    for r_id, r_name, r_desc in requirements:
        row_cells = table.add_row().cells
        row_cells[0].text = r_id
        row_cells[1].text = r_name
        row_cells[2].text = r_desc
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(51, 65, 85)

    add_heading_1("4. Casos de Uso")
    
    add_heading_2("UC01: Manter Produtos")
    add_body_text(" Usuário (Lojista)", "Ator:")
    add_body_text(" O usuário seleciona 'Produtos', informa o Código do produto. Se o produto não existe, o sistema pergunta se deseja cadastrar. Solicita Descrição, Categoria (0 a 5), Quantidade e Preço, e o código do Fornecedor. Caso o fornecedor não exista, inicia o fluxo UC02 de cadastro de fornecedor. Se o produto já existe, exibe os dados atuais e oferece opções para alterar ou excluir.", "Fluxo Principal:")
    
    add_heading_2("UC02: Manter Fornecedores")
    add_body_text(" Usuário (Lojista)", "Ator:")
    add_body_text(" O usuário seleciona 'Fornecedores', informa o Código. Se não existe, solicita CNPJ, Nome e Telefone. O CNPJ é verificado contra duplicidade. Se existe, exibe os dados e permite A/E/V.", "Fluxo Principal:")

    add_heading_2("UC03: Lançar Movimentação de Estoque")
    add_body_text(" Usuário (Lojista)", "Ator:")
    add_body_text(" O usuário acessa 'Movimentações', informa o Código da movimentação. Se não existe, solicita Data, Quantidade Movimentada e Código do Produto. O sistema busca o produto por código. Se o produto não for encontrado, oferece opção para cadastrar novo produto, listar cadastrados ou cancelar a operação. Se encontrado, valida se há quantidade em estoque suficiente. Se sim, abate a quantidade do estoque do produto e grava a movimentação.", "Fluxo Principal:")

    add_heading_2("UC04: Visualizar Relatórios")
    add_body_text(" Usuário (Lojista)", "Ator:")
    add_body_text(" O usuário seleciona o relatório correspondente (Produtos, Fornecedores ou Movimentações). O sistema limpa a tela, monta o cabeçalho e renderiza a lista formatada de registros. Ao fim, solicita pressionar qualquer tecla para retornar ao menu principal.", "Fluxo Principal:")

    add_heading_1("5. Classes e Estrutura do Código (Arquitetura do Projeto)")
    add_body_text("Abaixo estão detalhadas todas as classes geradas no código-fonte, divididas por camadas.")

    order = ['Models', 'Controllers', 'Views', 'Other']
    for category in order:
        if not code_info[category]:
            continue
        add_heading_2(f"Camada de {category}")
        for file_info in code_info[category]:
            for cls in file_info['classes']:
                base_desc = f" (Herda de: {cls['base']})" if cls['base'] else ""
                doc.add_paragraph().add_run(f"Classe: {cls['name']}{base_desc}").font.bold = True
                add_body_text(f"Arquivo: {file_info['filename']} | Namespace: {file_info['namespace']}")
                
                if cls['type'] == 'enum':
                    add_body_text("Valores do Enumerador:")
                    for f in cls['fields']:
                        if f['name'] not in ('value__', 'instance'):
                            add_bullet(f['name'])
                    continue

                if cls['fields'] or cls['properties']:
                    add_body_text("", "Membros (Atributos e Propriedades):")
                    c_table = doc.add_table(rows=1, cols=3)
                    c_table.style = 'Light Shading Accent 1'
                    c_hdr = c_table.rows[0].cells
                    c_hdr[0].text = 'Acesso'
                    c_hdr[1].text = 'Tipo'
                    c_hdr[2].text = 'Nome'
                    for cell in c_hdr:
                        set_cell_background(cell, "334155")
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        
                    for f in cls['fields']:
                        row = c_table.add_row().cells
                        row[0].text = f['access']
                        row[1].text = f['type']
                        row[2].text = f['name']
                    for p in cls['properties']:
                        row = c_table.add_row().cells
                        row[0].text = p['access']
                        row[1].text = p['type']
                        row[2].text = p['name'] + " { get; set; }"
                    
                    for r in c_table.rows[1:]:
                        for cell in r.cells:
                            cell.paragraphs[0].runs[0].font.size = Pt(8.5)
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(51, 65, 85)

                if cls['methods']:
                    add_body_text("", "Métodos:")
                    for m in cls['methods']:
                        mod = f"{m['modifier']} " if m['modifier'] else ""
                        args = f"({m['args']})" if m['args'] else "()"
                        add_code(f"{m['access']} {mod}{m['type']} {m['name']}{args}")
                doc.add_paragraph()

    add_heading_1("6. Mocks de Tela (ASCII Art)")
    
    add_heading_2("Tela Principal (Menu)")
    mock_menu = (
        "╔══════════════════════════════════════════════════════════════════════════════╗\n"
        "║                            Sistema StockShop                                 ║\n"
        "║                                                                              ║\n"
        "║   ╔══════════════════════════════════════════╗                               ║\n"
        "║   ║ 1 - Cadastrar Produto                    ║                               ║\n"
        "║   ║ 2 - Cadastrar Fornecedor                 ║                               ║\n"
        "║   ║ 3 - Movimentações de Produto             ║                               ║\n"
        "║   ║ 4 - Relatórios de Produtos Movimentados  ║                               ║\n"
        "║   ║ 5 - Relatórios de Produtos Cadastrados   ║                               ║\n"
        "║   ║ 6 - Relatórios de Fornecedores Cadastrados║                               ║\n"
        "║   ║ 0 - Sair                                 ║                               ║\n"
        "║   ║ Opção : _                                ║                               ║\n"
        "║   ╚══════════════════════════════════════════╝                               ║\n"
        "║                                                                              ║\n"
        "╚══════════════════════════════════════════════════════════════════════════════╝"
    )
    for line in mock_menu.split('\n'):
        add_code(line)

    add_heading_2("Tela de Cadastro de Produtos (CRUD)")
    mock_crud = (
        "╔══════════════════════════════════════════════════════════════════════════════╗\n"
        "║                            Cadastro de Produtos                              ║\n"
        "║                                                                              ║\n"
        "║   ╔══════════════════════════════════════════╗                               ║\n"
        "║   ║ Código                 : 1               ║                               ║\n"
        "║   ║ Descrição              : Lapis           ║                               ║\n"
        "║   ║ (0-Escolar, 1-Escritório,                ║                               ║\n"
        "║   ║  2-Presente, 3-Brinquedo,                ║                               ║\n"
        "║   ║  4-Artesanato, 5-Papelaria)              ║                               ║\n"
        "║   ║ Categoria              : 0               ║                               ║\n"
        "║   ║ Quantidade em Estoque  : 10              ║                               ║\n"
        "║   ║ Preço                  : R$2.50          ║                               ║\n"
        "║   ║ Código do Fornecedor   : 1               ║                               ║\n"
        "║   ╚══════════════════════════════════════════╝                               ║\n"
        "║                                                                              ║\n"
        "║   Confirma cadastro (S/N): _                                                 ║\n"
        "╚══════════════════════════════════════════════════════════════════════════════╝"
    )
    for line in mock_crud.split('\n'):
        add_code(line)

    add_heading_2("Tela de Relatório (Produtos Cadastrados)")
    mock_report = (
        "╔══════════════════════════════════════════════════════════════════════════════╗\n"
        "║                     Relatório de Produtos Cadastrados                        ║\n"
        "║                                                                              ║\n"
        "║  Código | Descrição           | Categoria     | Preço    | Estoque | Fornec. ║\n"
        "║  --------------------------------------------------------------------------- ║\n"
        "║  1      | Lapis               | Material_Esco | 2,50     | 10      | Art&Off ║\n"
        "║                                                                              ║\n"
        "║  Pressione qualquer tecla para voltar ao menu...                             ║\n"
        "╚══════════════════════════════════════════════════════════════════════════════╝"
    )
    for line in mock_report.split('\n'):
        add_code(line)

    doc.save(output_path)

if __name__ == '__main__':
    workspace = os.path.dirname(os.path.abspath(__file__))
    output_docx = os.path.join(workspace, "Documentacao_Sistema_StockShop.docx")
    build_docx(output_docx, workspace)
    print(f"DOCX gerado com sucesso em: {output_docx}")
